import json
import unittest
import uuid
from pathlib import Path

from docx import Document

from doc.build_top_contract_experiment_report import build_report


class BuildTopContractExperimentReportTests(unittest.TestCase):
    def test_report_includes_metric_generation_logic_explanations(self):
        payload = {
            'seed_contract': {
                'chain': 'ethereum',
                'contract_address': '0xseed',
                'token_type': 'ERC721',
                'contract_deployer': '0xcreator',
                'deployed_block_number': 123,
                'name': 'SeedProject',
                'symbol': 'SEED',
            },
            'seed_collection_stats': {
                'seed_nft_count': 2,
                'unique_token_uri_count': 2,
                'unique_image_uri_count': 2,
                'unique_name_count': 1,
                'unique_symbol_count': 1,
            },
            'suspected_infringing_duplicates_high_confidence': [
                {
                    'contract_address': '0xdup',
                    'candidate_count': 2,
                    'match_reasons': ['token_uri_match', 'image_uri_match'],
                }
            ],
            'suspected_infringing_duplicates_low_confidence': [],
            'legit_duplicates': [],
            'address_signals': {
                '0xdup': {
                    'mint_address_count': 2,
                    'mint_count': 3,
                    'unique_receiver_count': 4,
                    'cycle_edge_count': 1,
                    'star_distributor_count': 2,
                    'mint_to_first_transfer_seconds': 3200,
                    'fast_spread': True,
                }
            },
            'victim_signals': {
                '0xdup': {
                    'owner_count': 4,
                    'stuck_holder_count': 3,
                    'stuck_holder_ratio': 0.75,
                    'victim_wallet_count': 3,
                }
            },
            'report_summary': {
                'open_license_detected': False,
                'candidate_contract_count': 1,
                'high_confidence_contract_count': 1,
                'low_confidence_contract_count': 0,
                'legit_duplicate_contract_count': 0,
            },
        }

        tmp = Path('D:/code/solidity/get-nft-data/output/doc/test-tmp') / str(uuid.uuid4())
        tmp.mkdir(parents=True, exist_ok=True)
        try:
            summary_json = tmp / 'summary.json'
            summary_md = tmp / 'summary.md'
            source_docx = tmp / 'source.docx'
            output = tmp / 'report.docx'
            summary_json.write_text(json.dumps(payload, ensure_ascii=False), encoding='utf-8')
            summary_md.write_text('# summary', encoding='utf-8')
            Document().save(source_docx)

            build_report(summary_md=summary_md, summary_json=summary_json, source_docx=source_docx, output=output)

            text = '\n'.join(
                para.text for para in Document(output).paragraphs if para.text.strip()
            )
        finally:
            for path in sorted(tmp.glob('**/*'), reverse=True):
                if path.is_file():
                    path.unlink()
                elif path.is_dir():
                    path.rmdir()
            if tmp.exists():
                tmp.rmdir()

        self.assertIn('循环交易边数的统计逻辑', text)
        self.assertIn('星状扩散中心数的统计逻辑', text)
        self.assertIn('Mint 到首次转手时间的统计逻辑', text)
        self.assertIn('快速扩散字段的统计逻辑', text)


if __name__ == '__main__':
    unittest.main()
