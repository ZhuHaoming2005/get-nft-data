from __future__ import annotations

# Example:
#   C:\Users\z1766\.conda\envs\codex\python.exe D:\code\solidity\get-nft-data\doc\build_top_contract_experiment_report.py ^
#     --summary-md D:\code\solidity\get-nft-data\top_contract_analysis_2.md ^
#     --summary-json D:\code\solidity\get-nft-data\top_contract_analysis_2.json ^
#     --source-docx D:\code\solidity\get-nft-data\doc\NFT问题思路总结.docx ^
#     --output D:\code\solidity\get-nft-data\output\doc\Top_NFT_合约级重复样本实验报告.docx

import argparse
import json
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def _set_doc_defaults(doc: Document) -> None:
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    normal.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def _add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run(subtitle)
    run.font.size = Pt(11)
    run.italic = True
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_paragraph(doc: Document, text: str, bold_prefix: str = '') -> None:
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        text = text[len(bold_prefix):] if text.startswith(bold_prefix) else text
    p.add_run(text)


def _add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def _add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style='List Number')


def _format_ratio(value: float) -> str:
    return f'{value * 100:.2f}%'


def _load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding='utf-8'))


def _top_address_signals(address_signals: dict, limit: int = 3) -> list[tuple[str, dict]]:
    ranked = sorted(
        address_signals.items(),
        key=lambda item: (
            item[1].get('mint_to_first_transfer_seconds', 0),
            item[1].get('cycle_edge_count', 0),
            item[1].get('star_center_count', 0),
            item[1].get('mint_tx_count', 0),
        ),
        reverse=True,
    )
    return ranked[:limit]


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr[idx].text = text
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text


def build_report(summary_md: Path, summary_json: Path, source_docx: Path, output: Path) -> Path:
    payload = _load_json(summary_json)
    seed = payload.get('seed_contract') or {}
    seed_stats = payload.get('seed_collection_stats') or {}
    summary = payload.get('report_summary') or {}
    high = payload.get('suspected_infringing_duplicates_high_confidence') or []
    low = payload.get('suspected_infringing_duplicates_low_confidence') or []
    official = payload.get('legit_duplicates') or []
    address_signals = payload.get('address_signals') or {}
    victim_signals = payload.get('victim_signals') or {}

    doc = Document()
    _set_doc_defaults(doc)
    _add_title(
        doc,
        'Top NFT 合约级重复样本分析实验报告',
        '依据研究思路文档与当前 BAYC 合约级实验结果整理',
    )

    _add_heading(doc, '一、实验背景与目的', level=1)
    _add_paragraph(
        doc,
        '研究思路文档指出，NFT 版权问题与传统版权问题的核心差异在于链上 token 与链下作品相分离、'
        '智能合约执行缺乏事前版权审查、懒人铸造降低侵权成本，以及区块链匿名性和不可篡改性使侵权记录长期保留。'
        '在这种技术条件下，侵权 NFT 不仅表现为内容复制，还可能伴随洗售、虚假繁荣和散户套牢等市场伤害。',
    )
    _add_paragraph(
        doc,
        '因此，本实验的直接目标不是对单个 NFT 作法律裁判，而是以一个头部 NFT 合约为种子样本，'
        '识别其在链上被重复发行的外部合约，并进一步刻画这些重复样本是否呈现出异常传播和受害者沉淀。'
        '按照需求，这一目标可拆解为三层：第一，发现重复 NFT；第二，过滤出更接近侵权候选的样本；'
        '第三，围绕这些样本分析恶意地址和被骗地址的链上行为。',
    )
    _add_bullets(
        doc,
        [
            '实验对象：以以太坊链上的 BAYC 合约作为种子项目进行合约级重复样本分析。',
            '研究重点：从“头部 NFT 是否被复制”推进到“复制样本是否伴随异常传播与低流动性伤害”。',
            '当前版本边界：先基于链上数据完成查重、过滤和行为分析，价格和跨链扩展暂不纳入本次报告的实证部分。',
        ],
    )

    _add_heading(doc, '二、实验输入与数据来源', level=1)
    _add_bullets(
        doc,
        [
            f"种子链与合约：{seed.get('chain', '')} / {seed.get('contract_address', '')}",
            f"种子项目名称：{seed.get('name', '')}，符号：{seed.get('symbol', '')}，Token 类型：{seed.get('token_type', '')}",
            '外部 API：Alchemy 作为主数据源，用于拉取种子合约全部 NFT、合约元数据、单个 metadata、transfer 历史和 owner 快照；'
            'Etherscan 仅在 transfer 拉取失败时作为兜底。',
            '本地数据库：nft_assets_{chain} 作为查重主数据源，用于在已有 NFT 资产记录中查找重复样本。',
            f'结果文件参考：{summary_md.name} 与 {summary_json.name}',
            f'需求来源文档：{source_docx.name}',
        ],
    )

    _add_heading(doc, '三、实验流程与方法', level=1)
    _add_numbered(
        doc,
        [
            '种子集合生成：输入链与头部 NFT 合约地址后，先通过 Alchemy 获取该合约下全部 NFT，形成原始种子集合。'
            '当前实验中，BAYC 共拉取到 10000 个种子 NFT。',
            '查重键提取：从种子集合中提取 token_uri、image_uri、name、symbol 四类查重键。'
            '其中 token_uri 与 image_uri 为强证据，name 与 symbol 作为弱证据补充。',
            '本地数据库查重：将种子集合与本地 nft_assets_{chain} 进行比对。'
            '命中 token_uri 或 image_uri 的样本优先保留；仅命中 name 或 symbol 的样本会被降级处理。',
            '开放许可过滤：只对种子合约抽样拉取 1 个 metadata，检查是否含有 CC0-1.0、License: CC0、public domain 等文本。'
            '若命中开放许可，则不再输出侵权候选池。',
            '官方参与型重复过滤：对重复合约拉取 transfer 历史，提取所有 mint 接收地址；'
            '若 mint 接收地址与官方地址集合存在交集，则该合约被归入“官方参与型重复”分组。'
            '这里的官方地址集合当前由种子合约部署者和种子合约地址构成。',
            '恶意行为分析：对于高置信重复合约，继续分析 Mint 地址数、Mint 交易数、唯一接收地址数、循环交易边数、'
            '星状扩散中心数、Mint 到首次转手时间，以及是否出现快速扩散。',
            '受害者分析：结合 transfer 与 owner 快照，识别当前持有地址数、套牢地址数、套牢地址占比和疑似受害地址数，'
            '用以刻画重复样本的低流动性伤害。',
        ],
    )
    _add_paragraph(
        doc,
        '方法调整说明：研究思路文档原本提出第三类查重规则为 metadata JSON 文本近似匹配，但当前落地版本已改为'
        ' name 与 symbol 的近似/规范化匹配；开放许可也仅对原始合约抽样的 1 个 metadata 进行判断。'
        '因此，本次报告反映的是“当前已实现方法”的实验结果，而不是理想全量版本的全部能力。',
    )

    _add_heading(doc, '四、实验结果概述', level=1)
    _add_bullets(
        doc,
        [
            f"种子合约：{seed.get('name', '')}（{seed.get('symbol', '')}），合约地址 {seed.get('contract_address', '')}",
            f"部署者：{seed.get('contract_deployer', '')}，部署区块：{seed.get('deployed_block_number', 0)}",
            f"种子 NFT 数量：{seed_stats.get('seed_nft_count', 0)}，唯一 token URI 数：{seed_stats.get('unique_token_uri_count', 0)}，"
            f"唯一 image URI 数：{seed_stats.get('unique_image_uri_count', 0)}",
            f"开放许可检测结果：{'是' if summary.get('open_license_detected') else '否'}",
            f"重复候选合约数：{summary.get('candidate_contract_count', 0)}",
            f"高置信疑似侵权合约数：{summary.get('high_confidence_contract_count', 0)}",
            f"低置信疑似侵权合约数：{summary.get('low_confidence_contract_count', 0)}",
            f"被算法归为官方参与型重复的合约数：{summary.get('legit_duplicate_contract_count', 0)}",
        ],
    )
    _add_paragraph(
        doc,
        '从结果看，BAYC 作为头部 NFT 合约，在当前样本库中共对应 30 个重复候选合约，其中 29 个进入高置信疑似侵权分组，'
        '说明重复发行现象不是零星出现，而是以多合约、多样本的形式存在。'
        '更重要的是，高置信样本大多直接命中 image_uri 与 token_uri，这意味着它们不是简单借用相似名称，'
        '而是复用了与原始项目高度重合的资源标识。',
    )

    if high:
        rows = []
        for item in high[:10]:
            rows.append([
                item.get('contract_address', ''),
                str(item.get('candidate_count', 0)),
                ', '.join(item.get('match_reasons') or []),
            ])
        _add_paragraph(doc, '高置信疑似侵权合约示例（前 10 条）:')
        _add_table(doc, ['合约地址', '重复 NFT 数', '命中原因'], rows)

    if official:
        official_item = official[0]
        mint_recipients = official_item.get('mint_recipients') or []
        _add_heading(doc, '五、官方参与型重复样本解释', level=1)
        _add_paragraph(
            doc,
            '本次实验中只有 1 个合约被算法归入“官方参与型重复”分组，而没有进入高置信疑似侵权池。'
            '其原因不是已经获得法律上的授权证明，而是它满足当前实现中的官方地址交集规则。',
        )
        _add_bullets(
            doc,
            [
                f"合约地址：{official_item.get('contract_address', '')}",
                f"重复 NFT 数：{official_item.get('candidate_count', 0)}",
                f"规则依据：该合约的 mint 接收地址中出现了种子项目官方地址集合中的地址。",
                f"本次命中的关键官方地址：{seed.get('contract_deployer', '')}",
            ],
        )
        _add_paragraph(
            doc,
            '具体来说，当前代码会把“种子合约部署者地址”和“种子合约地址”组成官方地址集合；'
            '然后在重复合约的 transfer 历史中寻找 from=0x0 的 mint 记录，并收集其接收地址。'
            '只要这两类地址存在交集，该合约就被视为具有官方参与痕迹。'
            '因此，这里的“官方参与型重复”本质上是一种保守过滤标签，用于避免把可能与原项目直接相关的样本误判为外部侵权。'
            '它不等同于法律上已经确认授权，只表示当前链上证据显示官方地址曾参与该重复合约的初始接收环节。',
        )
        _add_paragraph(
            doc,
            '本次命中的 mint 接收地址较多，其中包含种子合约部署者地址，'
            f'因此该合约被归入官方参与型重复分组。mint 接收地址样例如下：{", ".join(mint_recipients[:6])}'
            + (' 等。' if len(mint_recipients) > 6 else '。'),
        )

    _add_heading(doc, '六、地址行为与受害者结果解释', level=1)
    top_signals = _top_address_signals(address_signals, limit=5)
    if top_signals:
        rows = []
        for contract_address, signal in top_signals:
            victim = victim_signals.get(contract_address) or {}
            rows.append([
                contract_address,
                str(signal.get('cycle_edge_count', 0)),
                str(signal.get('star_center_count', 0)),
                str(signal.get('mint_to_first_transfer_seconds', 0)),
                '是' if signal.get('fast_spread') else '否',
                str(victim.get('victim_wallet_count', 0)),
            ])
        _add_table(
            doc,
            ['合约地址', '循环交易边数', '星状扩散中心数', 'Mint 到首次转手时间(秒)', '快速扩散', '疑似受害地址数'],
            rows,
        )
    _add_paragraph(
        doc,
        '从行为层结果看，部分高置信重复合约已经表现出与研究设想一致的异常传播特征。'
        '例如部分合约出现较高的循环交易边数，说明少数地址之间可能存在反复倒手；'
        '另一些合约具有较多星状扩散中心，说明其传播并非自然形成，而更像面向大量外部地址的定向分发。'
        '在当前样本中，合约 0x82c7a8f707110f5fbb16184a5933e9f78a34c6ab 的 Mint 到首次转手时间为 3202 秒，'
        '并被识别为“快速扩散”，同时伴随 26 条循环交易边和 98 个星状扩散中心，是较典型的高风险传播样本。',
    )
    _add_paragraph(
        doc,
        '受害者层结果主要通过“当前仍持有且未出现卖出行为的地址”来刻画低流动性伤害。'
        '当前实现把这类地址定义为套牢地址，并进一步将其数量记录为疑似受害地址数。'
        '因此，当前版本已经能够部分回应“有多少买家被吸入这些重复样本”这一需求，'
        '但尚未进一步量化购买价格、资金规模和相对资产损失比例。',
    )

    _add_heading(doc, '七、核心指标生成逻辑说明', level=1)
    _add_paragraph(
        doc,
        '为避免把结果字段仅当作静态标签，本节说明几个核心行为指标在当前代码中的具体生成逻辑。'
        '这些逻辑直接对应 top_contract_analysis 模块中的 transfer 分析函数，因此属于当前实验的实际计算口径。',
    )
    _add_bullets(
        doc,
        [
            '循环交易边数的统计逻辑：先过滤掉 mint 和销毁相关的记录，即忽略 from=0x0 或 to=0x0 的 transfer；'
            '然后把每条普通转移记为一条有向边 (from_address, to_address)。如果后续出现反向边 (to_address, from_address)，'
            '则把这对地址记为一条循环交易边。最终统计的是发生过双向往返的地址对数量，而不是交易笔数。',
            '星状扩散中心数的统计逻辑：同样先排除 mint 和销毁记录，再统计每个发送地址发往多少个不同接收地址。'
            '若某个地址至少向 3 个不同接收者发送过 NFT，且它自己作为接收者被转入的次数不超过 1 次，'
            '则该地址被记为一个星状扩散中心。这个规则的目的是识别“单点向外大量分发、但自身很少被回流”的广播式传播者。',
            'Mint 到首次转手时间的统计逻辑：先找出所有 mint 记录，即 from_address=0x0 的 transfer，并取其中最早的 block_time 作为首次 mint 时间；'
            '再在所有非 mint 记录中取最早的 block_time 作为首次转手时间。两者的差值就是 Mint 到首次转手时间。'
            '这个字段完全来自 transfer API 的时间数据，不依赖本地查重库；当前版本已补上时间解析与区块时间回查逻辑。',
            '快速扩散字段的统计逻辑：当前实现把 Mint 到首次转手时间不为 0 且小于等于 24 小时的样本标记为“快速扩散”。'
            '也就是说，只有当重复合约在首次 mint 之后 24 小时内就出现了首笔非 mint 转手时，fast_spread 才为真。',
            '唯一接收地址数的统计逻辑：统计所有 transfer 中去重后的 to_address 数量，并排除零地址。'
            '它反映的是该合约样本在链上扩散到了多少个不同的钱包，而不是交易次数。',
            'Mint 地址数与 Mint 交易数的统计逻辑：Mint 地址数是所有 mint 记录的接收地址去重数量，Mint 交易数则是 mint 记录总条数。'
            '两者分别表示初始分发的覆盖面和初始发行强度。',
            '套牢地址数的统计逻辑：先用 owner 快照找出当前余额大于 0 的地址，再从 transfer 中识别所有出现过卖出行为的地址。'
            '当前仍持有 NFT、但从未作为卖方出现的地址，会被记为套牢地址；其数量和占比进一步形成疑似受害者指标。',
        ],
    )

    _add_heading(doc, '八、结果字段含义与需求对应', level=1)
    _add_paragraph(
        doc,
        '下表按照“输出字段 - 字段含义 - 对应需求”进行映射。'
        '其中“对应需求”中的编号，对应研究思路文档中的 2.1、2.2、2.3、2.4 四个实验环节。',
    )
    field_rows = [
        ['seed_contract.chain', '种子项目所在链', '实验输入标识', '总体输入'],
        ['seed_contract.contract_address', '头部 NFT 合约地址', '定义种子项目', '总体输入'],
        ['seed_contract.token_type', '种子合约类型，如 ERC721/1155', '辅助决定后续 transfer 拉取方式', '流程支持字段'],
        ['seed_contract.contract_deployer', '种子合约部署者地址', '构造官方地址集合，用于过滤官方参与型重复', '2.2'],
        ['seed_contract.deployed_block_number', '种子项目部署区块', '记录项目背景与链上起点', '辅助解释'],
        ['seed_contract.name / symbol', '项目名称与符号', '作为结果展示字段，也可参与弱匹配', '2.1'],
        ['seed_collection_stats.seed_nft_count', '种子合约 NFT 总数', '说明原始样本规模', '2.1'],
        ['seed_collection_stats.unique_token_uri_count', '种子集合中唯一 token_uri 数', '反映 token_uri 的唯一性与查重基础', '2.1'],
        ['seed_collection_stats.unique_image_uri_count', '种子集合中唯一 image_uri 数', '反映 image_uri 的唯一性与查重基础', '2.1'],
        ['seed_collection_stats.unique_name_count', '种子集合中规范化 name 的唯一值数量', '用于 name 弱匹配的覆盖程度', '2.1'],
        ['seed_collection_stats.unique_symbol_count', '种子集合中规范化 symbol 的唯一值数量', '用于 symbol 弱匹配的覆盖程度', '2.1'],
        ['duplicate_candidates', '所有重复候选 NFT 明细', '查重后的原始候选池', '2.1'],
        ['duplicate_candidates[].match_reasons', '命中原因，如 token_uri_match、image_uri_match、name_match、symbol_match', '说明重复证据来自哪一层', '2.1'],
        ['duplicate_candidates[].confidence', '高/低置信标签', '决定候选样本是否进入重点分析', '2.2'],
        ['suspected_infringing_duplicates_high_confidence', '高置信疑似侵权合约聚合结果', '后续恶意地址与受害者分析的核心对象', '2.2 / 2.3 / 2.4'],
        ['suspected_infringing_duplicates_low_confidence', '低置信疑似侵权合约聚合结果', '保留弱证据样本，避免误判', '2.2'],
        ['legit_duplicates', '被算法归为官方参与型重复的合约', '排除可能与原项目存在直接链上联系的样本', '2.2'],
        ['report_summary.open_license_detected', '是否检测到开放许可', '若为真，则不进入侵权候选分析', '2.2'],
        ['report_summary.candidate_contract_count', '重复候选合约数', '刻画复制现象的范围', '2.1'],
        ['report_summary.high_confidence_contract_count', '高置信疑似侵权合约数', '刻画高风险复制范围', '2.2'],
        ['report_summary.low_confidence_contract_count', '低置信疑似侵权合约数', '刻画弱证据复制范围', '2.2'],
        ['report_summary.legit_duplicate_contract_count', '被算法归为官方参与型重复的合约数', '表示被过滤出的官方参与样本规模', '2.2'],
        ['address_signals[].mint_address_count', '参与 mint 接收的地址数', '衡量发行与初始分发范围', '2.3'],
        ['address_signals[].mint_tx_count', 'mint 交易数', '衡量侵权样本初始发行强度', '2.3'],
        ['address_signals[].unique_receiver_count', '唯一接收地址数', '观察扩散范围', '2.3'],
        ['address_signals[].cycle_edge_count', '循环交易边数', '识别关联地址高频互转', '2.3'],
        ['address_signals[].star_center_count', '星状扩散中心数', '识别单点向大量地址单向扩散', '2.3'],
        ['address_signals[].mint_to_first_transfer_seconds', '从首次 mint 到首次非 mint 转手的时间差', '衡量传播爆发速度', '2.3'],
        ['address_signals[].fast_spread', '是否快速扩散', '把时间指标转成易读风险标记', '2.3'],
        ['victim_signals[].owner_count', '当前持有该合约 NFT 的地址数', '估计当前暴露面', '2.4'],
        ['victim_signals[].stuck_holder_count', '当前持有且未出现卖出行为的地址数', '刻画被套牢的钱包规模', '2.4'],
        ['victim_signals[].stuck_holder_ratio', '套牢地址占当前持有地址的比例', '刻画流动性伤害深度', '2.4'],
        ['victim_signals[].victim_wallet_count', '当前版本中的疑似受害地址数', '作为被骗地址规模的近似指标', '2.4'],
    ]
    _add_table(doc, ['字段', '字段含义', '在当前实验中的作用', '对应需求'], field_rows)

    _add_heading(doc, '九、字段与需求的整体对应关系总结', level=1)
    _add_bullets(
        doc,
        [
            '需求 2.1“筛选所有重复的 NFT，记录对应地址”主要由 duplicate_candidates、seed_collection_stats、report_summary.candidate_contract_count 支撑。',
            '需求 2.2“筛选出侵犯版权的 NFT 并记录历史”在当前版本中对应高/低置信疑似侵权合约、开放许可过滤和官方参与型重复过滤。'
            '需要注意，这一步当前仍属于算法上的侵权候选筛选，不是法律意义上的最终认定。',
            '需求 2.3“识别恶意地址与被骗地址”主要由 address_signals 承担，其中循环交易、星状扩散和传播时间是恶意行为的核心代理变量。',
            '需求 2.4“分析被骗交易记录”在当前版本中被部分实现，主要落在 victim_signals 上，能够刻画地址规模与套牢程度；'
            '但价格、出售记录、资金损失比例等更强指标还需要补充市场交易金额数据。',
            '需求 2.5“跨链扩展”尚未进入这次实证报告，当前报告仅覆盖单链、单种子合约的实验闭环。',
        ],
    )

    _add_heading(doc, '十、结论与当前局限', level=1)
    _add_paragraph(
        doc,
        '本次实验已经完成了“头部 NFT 合约 -> 重复样本发现 -> 候选过滤 -> 异常传播/受害者分析”的单链闭环。'
        '以 BAYC 为例，结果表明该项目在本地样本库中存在成规模的外部重复合约，其中大多数具有较强的资源重合证据，'
        '并且部分样本伴随明显的异常传播特征和疑似受害者沉淀。这说明 NFT 版权问题不仅体现为内容复制，'
        '还可能进一步转化为面向购买者的市场伤害。',
    )
    _add_paragraph(
        doc,
        '与此同时，当前版本仍有三个边界。第一，官方参与型重复的判断依赖“mint 接收地址与官方地址交集”这一启发式规则，'
        '因此只能视为过滤标签，而不能直接当作法律授权证明。第二，当前未纳入价格和交易金额，'
        '因此无法对“套牢资金规模”和“买入金额占资产比例”作出完整量化。第三，跨链扩展尚未实做，'
        '因此本报告的结论只适用于当前单链实验结果。',
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return output


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description='Build a DOCX experiment report for top NFT contract duplicate analysis.')
    parser.add_argument('--summary-md', required=True, help='path to rendered markdown summary')
    parser.add_argument('--summary-json', required=True, help='path to analysis json payload')
    parser.add_argument('--source-docx', required=True, help='path to the requirement / experiment idea docx')
    parser.add_argument('--output', required=True, help='path to output docx')
    return parser


def main() -> int:
    args = build_parser().parse_args()
    output = build_report(
        summary_md=Path(args.summary_md),
        summary_json=Path(args.summary_json),
        source_docx=Path(args.source_docx),
        output=Path(args.output),
    )
    print(output)
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
